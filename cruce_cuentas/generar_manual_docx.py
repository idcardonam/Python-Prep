# -*- coding: utf-8 -*-
"""Genera MANUAL_OPERACION.docx (Word) para el área."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).resolve().parent / "MANUAL_OPERACION.docx"
AZUL = RGBColor(0x1B, 0x3A, 0x5C)
DORADO = RGBColor(0xC4, 0xA0, 0x35)
GRIS = RGBColor(0x4A, 0x55, 0x68)
NEGRO = RGBColor(0x1A, 0x20, 0x2C)


def set_run_font(run, name="Calibri", size=11, bold=False, italic=False, color=NEGRO):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def add_p(doc, text, *, size=11, bold=False, italic=False, color=NEGRO, align="left", space_after=8, space_before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return p


def add_h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, name="Calibri", size=16 if level == 1 else 13, bold=True, color=AZUL)
    p.paragraph_format.space_before = Pt(16 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_table(doc, rows: list[list[str]], widths=None):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, fila in enumerate(rows):
        for j, val in enumerate(fila):
            cell = t.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            if i == 0:
                set_run_font(run, size=9, bold=True, color=RGBColor(255, 255, 255))
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "1B3A5C")
                shd.set(qn("w:val"), "clear")
                tcPr.append(shd)
            else:
                set_run_font(run, size=9, color=NEGRO)
                if i % 2 == 0:
                    tcPr = cell._tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:fill"), "F7FAFC")
                    shd.set(qn("w:val"), "clear")
                    tcPr.append(shd)
    if widths:
        for row in t.rows:
            for idx, w in enumerate(widths):
                row.cells[idx].width = Cm(w)
    doc.add_paragraph()
    return t


def construir() -> Path:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = header.add_run("UNAB  ·  Dirección de TIC  ·  Estadísticas TIC  ·  Portal de cuentas Gmail")
        set_run_font(r, size=9, color=AZUL)
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = footer.add_run("Uso interno del área  ·  No publicar CSV de origen")
        set_run_font(r, size=8, italic=True, color=GRIS)

    add_p(doc, "UNIVERSIDAD AUTÓNOMA DE BUCARAMANGA", size=10, bold=True, color=DORADO, align="center", space_after=2)
    add_p(
        doc,
        "Dirección de Tecnologías de la Información y las Comunicaciones",
        size=11,
        italic=True,
        color=GRIS,
        align="center",
        space_after=10,
    )
    add_p(doc, "Manual de operación", size=22, bold=True, color=AZUL, align="center", space_after=4)
    add_p(
        doc,
        "Portal de gestión de cuentas Gmail institucionales",
        size=14,
        italic=True,
        color=GRIS,
        align="center",
        space_after=4,
    )
    add_p(
        doc,
        "Depuración de cuentas · Campaña 2FA · Seguimiento en SharePoint",
        size=11,
        color=NEGRO,
        align="center",
        space_after=4,
    )
    add_p(
        doc,
        "Versión septiembre 2026  ·  Documento para uso general del área",
        size=10,
        italic=True,
        color=GRIS,
        align="center",
        space_after=16,
    )

    add_h(doc, "1. Propósito", 1)
    add_p(
        doc,
        "Este portal reúne en un solo enlace el seguimiento de cuentas Gmail institucionales: "
        "depuración de cuentas (inventario, bloqueos, inactividad); campaña de autenticación en dos pasos "
        "(2FA) de estudiantes vigentes; y registro de la meta del proyecto y de las acciones realizadas. "
        "No es necesario saber programar para consultar el portal ni para la actualización periódica de archivos.",
        align="justify",
    )
    add_p(doc, "Dirección del portal", bold=True, size=11, space_after=2)
    add_p(
        doc,
        "https://unabedu.sharepoint.com/sites/ProyectoDepuracinGmail/SitePages/CollabHome.aspx",
        size=10,
        color=AZUL,
    )
    add_p(
        doc,
        "El sitio SharePoint se llama Proyecto Depuración Gmail. Es un grupo privado: solo entra quien tenga permiso explícito.",
        align="justify",
    )

    add_h(doc, "2. Qué hay en el portal (tres piezas)", 1)
    add_table(
        doc,
        [
            ["Pieza", "Qué se ve", "Cómo se mantiene"],
            [
                "Tablero Power BI",
                "Totales, estados de cuenta, avance",
                "Tres CSV en etl/output y listas. Power BI se actualiza cada hora.",
            ],
            [
                "Resumen 2FA",
                "Cobertura por facultad y pendientes",
                "Informe HTML y un CSV de pendientes en etl/output.",
            ],
            [
                "Listas SharePoint",
                "Meta del proyecto y bitácora de acciones",
                "Se editan en el sitio, a mano. No las genera el actualizador.",
            ],
        ],
        [4.2, 5.5, 6.5],
    )
    add_p(
        doc,
        "Las listas no son archivos. No se «suben». Se abren en el menú izquierdo (MetaProyecto, Acciones) "
        "y se editan como una tabla. Los números del inventario no salen de las listas Cuentas / Dependencias / capacidad. "
        "Salen de los CSV cuentas_powerbi.csv, dependencias_powerbi.csv y capacidad_powerbi.csv.",
        align="justify",
    )

    add_h(doc, "3. Quién hace qué y cómo se da acceso", 1)
    add_p(
        doc,
        "Hay dos perfiles. Se asignan en SharePoint. Si el informe está incrustado en el sitio, en muchos casos basta el permiso del sitio; si el tablero pide licencia, se da acceso también en Power BI (apartado 3.3).",
        align="justify",
    )

    add_h(doc, "3.1 Consulta (revisar)", 2)
    add_p(
        doc,
        "Puede abrir el portal, ver el tablero, el Resumen 2FA y las listas. No reemplaza archivos en etl/output.",
        align="justify",
    )
    add_p(doc, "Cómo dar este acceso", bold=True, space_after=4)
    for linea in [
        "1. Entre al sitio Proyecto Depuración Gmail.",
        "2. Arriba a la derecha: engranaje → Permisos del sitio (o Información del sitio → Permisos).",
        "3. Compartir sitio (o Invitar a personas).",
        "4. Escriba el correo institucional de la persona.",
        "5. Nivel: Lectura (o grupo Visitantes).",
        "6. Confirme. La persona recibe correo o verá el sitio en SharePoint.",
    ]:
        add_p(doc, linea, space_after=3)
    add_p(doc, "También puede usar Compartir en la página de inicio y elegir «Puede ver».", align="justify")

    add_h(doc, "3.2 Operación (cargar archivos y editar listas)", 2)
    add_p(
        doc,
        "Puede subir los seis archivos a etl/output, editar MetaProyecto y Acciones, y usar actualizar.bat en el equipo del área.",
        align="justify",
    )
    add_p(doc, "Cómo dar este acceso: mismos pasos 1 a 4 de consulta, con nivel Edición (grupo Miembros).", align="justify")
    add_p(
        doc,
        "Quien opera no debe dejar etl/input abierto a todo el sitio: esa carpeta guarda CSV de origen. "
        "En la carpeta input: … → Administrar acceso → dejar solo al personal de operación.",
        align="justify",
    )

    add_h(doc, "3.3 Power BI", 2)
    add_p(doc, "Si al abrir el tablero pide licencia o «no tiene permiso»:", align="justify", space_after=4)
    for linea in [
        "1. En app.powerbi.com abra el área de trabajo del informe.",
        "2. Acceso (o Compartir informe).",
        "3. Agregue el mismo correo, rol Visor (consulta) o Colaborador (si debe refrescar el conjunto de datos).",
    ]:
        add_p(doc, linea, space_after=3)

    add_h(doc, "4. Cómo consultar (uso diario)", 1)
    for linea in [
        "1. Abra el enlace del portal (sección 1). Inicie sesión con la cuenta UNAB.",
        "2. En la portada verá el tablero de depuración.",
        "3. Resumen 2FA: abre el informe de estudiantes vigentes sin 2FA. Use facultad y programa. Descargar CSV pendientes abre un Excel; filtre la columna facultad.",
        "4. MetaProyecto: consulte o ajuste la meta (solo con permiso de edición).",
        "5. Acciones: consulte o registre lo ejecutado (bloqueos, comunicados, cierres de lote).",
    ]:
        add_p(doc, linea, space_after=3)
    add_p(
        doc,
        "Si el tablero se ve desactualizado, espere el refresco horario o pida a quien opera «Actualizar ahora» en Power BI.",
        align="justify",
    )

    add_h(doc, "5. Actualización periódica de datos (quien opera)", 1)
    add_p(
        doc,
        "Frecuencia sugerida: cada vez que haya un corte nuevo de Google Admin (típicamente mensual).",
        align="justify",
    )

    add_h(doc, "5.1 Qué reunir en UNA carpeta del PC (entrada)", 2)
    add_p(
        doc,
        "Solo insumos de sistemas. No mezcle aquí los HTML ni los CSV *_powerbi.",
        align="justify",
    )
    add_table(
        doc,
        [
            ["Archivo", "Obligatorio", "Origen"],
            [
                "User_Download_….csv",
                "Sí",
                "Google Admin Console → informe de usuarios (todas las cuentas del dominio).",
            ],
            [
                "CSV de inscritos o prematriculados (uno o varios)",
                "Sí (para 2FA)",
                "Extracto académico del periodo.",
            ],
            ["VISTA DE CURRICULO.xlsx", "No", "Catálogo de planes."],
        ],
        [5.0, 3.2, 8.0],
    )
    add_p(doc, "El mismo User_Download sirve para el tablero de depuración y para el cruce 2FA.", align="justify")

    add_h(doc, "5.2 Ejecutar el actualizador", 2)
    for linea in [
        "1. En el equipo del área: doble clic en Python-Prep\\cruce_cuentas\\actualizar.bat.",
        "2. Cuando pida la ruta: en el Explorador abra la carpeta del 5.1, clic en la barra de dirección, copie (Ctrl+C) y pegue en la ventana azul. Enter.",
        "3. Espere 1 a 3 minutos. Debe aparecer LISTO.",
        "4. Se abre una carpeta en el escritorio: Archivos_SharePoint_AAAA-MM-DD (la fecha del día evita mezclar cortes).",
    ]:
        add_p(doc, linea, space_after=3)

    add_h(doc, "5.3 Qué hay que cargar en SharePoint (solo seis archivos)", 2)
    add_table(
        doc,
        [
            ["Archivo", "Función"],
            ["cuentas_powerbi.csv", "Tablero: detalle de cuentas"],
            ["dependencias_powerbi.csv", "Tablero: resumen por dependencia"],
            ["capacidad_powerbi.csv", "Tablero: licencias e inventario"],
            ["resumen.html", "Informe 2FA para consulta"],
            ["listado_sin_2fa.html", "Listado 2FA de trabajo"],
            ["02_estudiantes_sin_2fa.csv", "Descarga de pendientes (filtrar facultad en Excel)"],
        ],
        [7.0, 9.2],
    )
    add_p(
        doc,
        "No suba el archivo LEAME, ni los sin_2fa_….csv por facultad, ni 00_universo.csv. Esos quedan en el PC, en cruce_cuentas\\salida\\.",
        align="justify",
    )

    add_h(doc, "5.4 Cómo cargar", 2)
    for linea in [
        "1. En el sitio: Documentos → carpeta etl → output.",
        "2. Seleccione los seis archivos de la carpeta del escritorio.",
        "3. Arrástrelos a output.",
        "4. Si pregunta ¿Reemplazar?, elija Reemplazar.",
        "5. Abra el portal y compruebe el Resumen 2FA (fecha del corte).",
        "6. Power BI: se verá el corte en la próxima hora, o Actualizar ahora en el conjunto de datos (app.powerbi.com).",
    ]:
        add_p(doc, linea, space_after=3)

    add_h(doc, "5.5 Power BI — programación cada hora (una sola vez)", 2)
    for linea in [
        "1. Entre a https://app.powerbi.com con cuenta UNAB.",
        "2. Área de trabajo del proyecto → el conjunto de datos (no el informe).",
        "3. … → Configuración → Actualización programada.",
        "4. Activar. Frecuencia: Cada hora. Zona: Bogotá.",
        "5. Guardar.",
    ]:
        add_p(doc, linea, space_after=3)
    add_p(
        doc,
        "No es necesario volver a publicar el archivo .pbix si las fuentes ya apuntan a SharePoint.",
        align="justify",
    )

    add_h(doc, "5.6 Error Premium_ASWL_Error / Workspace Identity (el tablero no refresca)", 2)
    add_p(
        doc,
        "Este error no se debe a los CSV. El conjunto de datos está autenticado con «identidad del área de trabajo» "
        "(Workspace Identity) y esa identidad no existe, o quien publica el modelo no tiene permiso de Colaborador "
        "(o superior) en el área de trabajo de Power BI.",
        align="justify",
    )
    add_p(doc, "Camino rápido (recomendado para este portal): cuenta organizacional, no identidad del área.", bold=True)
    for linea in [
        "1. Entre a https://app.powerbi.com",
        "2. Área de trabajo del informe → el conjunto de datos (modelo semántico) → … → Configuración.",
        "3. Abra Credenciales del origen de datos.",
        "4. En cada origen SharePoint (carpeta etl/output y listas): Editar credenciales.",
        "5. Método de autenticación: OAuth2 o Cuenta organizacional (no «Identidad del área de trabajo»).",
        "6. Inicie sesión con una cuenta institucional que sí tenga acceso al sitio Proyecto Depuración Gmail.",
        "7. Nivel de privacidad: Organizacional.",
        "8. Guardar → Actualizar ahora.",
    ]:
        add_p(doc, linea, space_after=3)
    add_p(doc, "Camino alternativo: crear la identidad del área de trabajo.", bold=True)
    for linea in [
        "1. Área de trabajo → Configuración del área de trabajo → Identidad del área de trabajo → Crear.",
        "2. Dé a esa identidad acceso al sitio SharePoint (al menos lectura en etl/output y listas).",
        "3. Quien es propietario del modelo debe ser Colaborador o superior en el área de trabajo de Power BI.",
        "4. Actualizar ahora.",
    ]:
        add_p(doc, linea, space_after=3)
    add_p(
        doc,
        "Si el área de trabajo no es de capacidad Fabric/Premium, use el primer camino (OAuth2).",
        align="justify",
    )

    add_h(doc, "6. Listas: MetaProyecto y Acciones", 1)
    add_table(
        doc,
        [
            ["Lista", "Para qué", "Quién la cambia"],
            ["MetaProyecto", "Meta de depuración, fechas, proyección", "Quien define la meta del periodo"],
            ["Acciones", "Bitácora (qué se hizo, cuándo)", "Quien ejecuta la operación"],
        ],
        [4.0, 6.5, 5.7],
    )
    add_p(
        doc,
        "Para editar: menú izquierdo → nombre de la lista → Editar o + Nuevo. En minutos u hora, Power BI lo toma. "
        "Si cambia la meta y el tablero no se mueve, no regenere CSV: edite la lista y refresque Power BI.",
        align="justify",
    )

    add_h(doc, "7. Flujo resumido", 1)
    add_p(
        doc,
        "Carpeta del corte (Google Admin + inscritos) → doble clic en actualizar.bat → "
        "Escritorio\\Archivos_SharePoint_fecha (6 archivos) → arrastrar a SharePoint etl/output → "
        "tablero (cada hora) y botón Resumen 2FA. Listas MetaProyecto / Acciones se editan en el sitio y entran en el mismo refresco horario.",
        align="justify",
    )

    add_h(doc, "8. Incidencias frecuentes", 1)
    add_table(
        doc,
        [
            ["Qué ocurre", "Qué hacer"],
            ["No se encontró Python", "Instalar Python desde python.org y marcar Add Python to PATH."],
            ["No hallé export Google", "El archivo debe ser User_Download… y estar en la carpeta que pegó."],
            ["No hallé CSV de inscritos", "Faltan los archivos académicos en esa misma carpeta."],
            ["El tablero no cambia", "Confirme que reemplazó los tres *_powerbi.csv. Pulse Actualizar ahora."],
            [
                "Error Premium_ASWL / Workspace Identity",
                "Apartado 5.6: cambiar credenciales a cuenta organizacional (OAuth2).",
            ],
            ["La meta no cambia", "Edite MetaProyecto, no un CSV."],
            ["404 al descargar pendientes", "Falta 02_estudiantes_sin_2fa.csv en etl/output, junto al HTML."],
            ["No aparece la carpeta en el escritorio", "Busque Escritorio o Desktop. El actualizador la abre al terminar."],
            ["No tiene acceso al sitio", "Pedir inclusión con lectura o edición (sección 3)."],
            ["El HTML 2FA se ve raro", "Ábralo en pestaña nueva, no solo en la vista previa de SharePoint."],
        ],
        [5.8, 10.4],
    )

    add_h(doc, "9. Datos personales", 1)
    add_p(
        doc,
        "Los CSV de origen y el universo de cuentas contienen información sensible. Trabaje en el equipo del área. "
        "No publique etl/input ni 00_universo.csv en el portal de consulta. No envíe listados de correos por canales no institucionales.",
        align="justify",
    )

    add_h(doc, "10. Soporte", 1)
    add_p(doc, "Estadísticas TIC — Dirección de TIC", bold=True, space_after=2)
    add_p(doc, "Universidad Autónoma de Bucaramanga")

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(construir())
